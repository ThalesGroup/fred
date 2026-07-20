# Copyright Thales 2026
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

"""Minimal Markdown -> Word (.docx) conversion for writable documents (#1905 port).

Intentionally pragmatic (v1): supports the constructs an agent typically emits for
a deliverable — headings, paragraphs, bold/italic, inline code, bullet/numbered
lists, fenced code blocks and GFM tables. It is not a full CommonMark implementation.

Ported near-verbatim from Kea's ``core/writable_documents/docx_export.py``; the inline
grammar lives once in the shared :mod:`fred_capability_writable_document.inline` parser.
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Pt
from docx.text.paragraph import Paragraph

from fred_capability_writable_document.inline import parse_inline_markdown

_HEADING_RE = re.compile(r"^(?P<hashes>#{1,6})\s+(?P<text>.*)$")
_BULLET_RE = re.compile(r"^(?P<indent>\s*)[-*+]\s+(?P<text>.*)$")
_ORDERED_RE = re.compile(r"^(?P<indent>\s*)\d+[.)]\s+(?P<text>.*)$")
_FENCE_RE = re.compile(r"^\s*```")

# List nesting: every 2 columns of leading whitespace is one indent level (a tab
# counts as 2 columns). python-docx ships styles up to level 3 ("List Bullet",
# "List Bullet 2", "List Bullet 3"); deeper nesting is capped at level 3.
_LIST_INDENT_WIDTH = 2
_MAX_LIST_LEVEL = 3


def _list_style(base: str, indent: str) -> str:
    """Map a list item's leading whitespace to a python-docx list style name.

    ``base`` is ``"List Bullet"`` or ``"List Number"``. Level 1 keeps the base
    name; levels 2..3 append the number (e.g. ``"List Bullet 2"``).
    """
    columns = indent.replace("\t", " " * _LIST_INDENT_WIDTH).count(" ")
    level = min(columns // _LIST_INDENT_WIDTH + 1, _MAX_LIST_LEVEL)
    return base if level == 1 else f"{base} {level}"


# A table row contains at least one (unescaped) pipe. The separator line that
# follows the header is made only of pipes, dashes, colons and whitespace, with
# at least one dash — e.g. `|---|:--:|`. The separator is what distinguishes a
# real table from a paragraph that happens to contain a literal `|`.
_TABLE_ROW_RE = re.compile(r"^\s*\|?.*\|.*$")
_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)*\|?\s*$")


def _add_inline_runs(paragraph: Paragraph, text: str) -> None:
    """Render a line of Markdown inline formatting into runs on a paragraph.

    Thin consumer of the shared :func:`parse_inline_markdown`: it owns only the docx
    step of writing each parsed span as a run, toggling ``.bold`` / ``.italic`` and
    swapping a code span to Courier. The grammar lives once in the shared parser.
    """
    for span in parse_inline_markdown(text):
        run = paragraph.add_run(span.text)
        if span.bold:
            run.bold = True
        if span.italic:
            run.italic = True
        if span.code:
            run.font.name = "Courier New"


def _split_table_row(line: str) -> list[str]:
    """Split a Markdown table row into cell texts, honoring escaped pipes (\\|)."""
    # Replace escaped pipes with a placeholder, split on real pipes, restore.
    placeholder = "\x00"
    cells = line.replace(r"\|", placeholder).split("|")
    # Drop the empty leading/trailing cells produced by surrounding pipes.
    if cells and cells[0].strip() == "":
        cells = cells[1:]
    if cells and cells[-1].strip() == "":
        cells = cells[:-1]
    return [c.replace(placeholder, "|").strip() for c in cells]


def _is_table_start(lines: list[str], i: int) -> bool:
    """A table is a row line immediately followed by a separator line."""
    return (
        i + 1 < len(lines)
        and _TABLE_ROW_RE.match(lines[i]) is not None
        and "|" in lines[i]
        and _TABLE_SEP_RE.match(lines[i + 1]) is not None
    )


def _add_table(document: DocxDocument, rows: list[list[str]]) -> None:
    """Render parsed table rows (first row = header) into a styled docx table."""
    n_cols = max(len(r) for r in rows)
    table = document.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            # add_table seeds each cell with an empty paragraph; reuse it.
            paragraph = cell.paragraphs[0]
            _add_inline_runs(paragraph, text)
            if r_idx == 0:
                for run in paragraph.runs:
                    run.bold = True


def _add_code_block(document: DocxDocument, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def markdown_to_docx_bytes(content_md: str, *, title: str | None = None) -> bytes:
    """Convert Markdown text to a .docx document and return it as bytes."""
    document = Document()
    if title:
        _add_inline_runs(document.add_heading(level=0), title)

    lines = (content_md or "").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Fenced code block: consume until the closing fence.
        if _FENCE_RE.match(line):
            block: list[str] = []
            i += 1
            while i < len(lines) and not _FENCE_RE.match(lines[i]):
                block.append(lines[i])
                i += 1
            i += 1  # skip closing fence (or EOF)
            _add_code_block(document, block)
            continue

        # Blank line -> paragraph break (skip; docx paragraphs separate blocks).
        if not line.strip():
            i += 1
            continue

        heading = _HEADING_RE.match(line)
        if heading:
            level = len(heading.group("hashes"))
            paragraph = document.add_heading(level=min(level, 6))
            _add_inline_runs(paragraph, heading.group("text").strip())
            i += 1
            continue

        # Table: a row line followed by a separator line. Consume contiguous rows.
        if _is_table_start(lines, i):
            header = _split_table_row(lines[i])
            i += 2  # skip header and separator
            body = []
            while i < len(lines) and _TABLE_ROW_RE.match(lines[i]) and "|" in lines[i]:
                body.append(_split_table_row(lines[i]))
                i += 1
            _add_table(document, [header, *body])
            continue

        bullet = _BULLET_RE.match(line)
        if bullet:
            style = _list_style("List Bullet", bullet.group("indent"))
            paragraph = document.add_paragraph(style=style)
            _add_inline_runs(paragraph, bullet.group("text").strip())
            i += 1
            continue

        ordered = _ORDERED_RE.match(line)
        if ordered:
            style = _list_style("List Number", ordered.group("indent"))
            paragraph = document.add_paragraph(style=style)
            _add_inline_runs(paragraph, ordered.group("text").strip())
            i += 1
            continue

        # Default: a normal paragraph. Merge consecutive plain lines into one
        # paragraph (soft-wrapped Markdown), stopping at blanks/structural lines.
        para_lines = [line.strip()]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if (
                not nxt.strip()
                or _HEADING_RE.match(nxt)
                or _BULLET_RE.match(nxt)
                or _ORDERED_RE.match(nxt)
                or _FENCE_RE.match(nxt)
            ):
                break
            para_lines.append(nxt.strip())
            i += 1
        paragraph = document.add_paragraph()
        _add_inline_runs(paragraph, " ".join(para_lines))

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()

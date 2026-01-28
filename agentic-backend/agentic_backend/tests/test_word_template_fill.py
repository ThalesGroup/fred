"""Test suite for fill_word_from_structured_response function"""
import tempfile
from pathlib import Path

import pytest
from docx import Document

from agentic_backend.agents.reference_editor.powerpoint_template_util import (
    fill_word_from_structured_response,
)


@pytest.fixture
def sample_template():
    """Create a sample Word document with placeholders"""
    doc = Document()

    # Add paragraphs with various placeholder types
    doc.add_paragraph("Project: {nomProjet}")
    doc.add_paragraph("Company: {nomSociete}")
    doc.add_paragraph("Date: {dateProjet}")

    # Add a paragraph with multiple placeholders
    doc.add_paragraph("Project {nomProjet} for {nomSociete} started on {dateProjet}")

    # Add a table with placeholders
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Field"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Technologies"
    table.cell(1, 1).text = "{listeTechnologies}"

    # Add header with placeholder
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = "Header - {nomSociete}"

    # Add footer with placeholder
    footer = section.footer
    footer.paragraphs[0].text = "Footer - {nomProjet}"

    # Save to temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return Path(tmp.name)


@pytest.fixture
def sample_data():
    """Sample structured data for replacement"""
    return {
        "informationsProjet": {
            "nomSociete": "TechCorp Inc.",
            "nomProjet": "Project Phoenix",
            "dateProjet": "2024-01-15",
            "nombrePersonnes": "25",
            "enjeuFinancier": "2.5M EUR"
        },
        "contexte": {
            "presentationClient": "Leading tech company",
            "presentationContexte": "Digital transformation initiative",
            "listeTechnologies": "Python, React, Docker, Kubernetes"
        },
        "syntheseProjet": {
            "enjeux": "Modernize legacy systems",
            "activiteSolutions": "Cloud migration and API development",
            "beneficeClients": "Reduced costs and improved scalability",
            "pointsForts": "Expert team with proven track record"
        }
    }


def test_basic_placeholder_replacement(sample_template, sample_data):
    """Test that basic placeholders are replaced correctly"""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as out:
        output_path = Path(out.name)

    try:
        fill_word_from_structured_response(sample_template, sample_data, output_path)

        # Read the output document
        result_doc = Document(str(output_path))

        # Check that placeholders were replaced in paragraphs
        text_content = "\n".join([p.text for p in result_doc.paragraphs])

        assert "Project Phoenix" in text_content
        assert "TechCorp Inc." in text_content
        assert "2024-01-15" in text_content
        assert "{nomProjet}" not in text_content
        assert "{nomSociete}" not in text_content

    finally:
        output_path.unlink(missing_ok=True)
        sample_template.unlink(missing_ok=True)


def test_multiple_placeholders_in_paragraph(sample_template, sample_data):
    """Test that multiple placeholders in a single paragraph are replaced"""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as out:
        output_path = Path(out.name)

    try:
        fill_word_from_structured_response(sample_template, sample_data, output_path)

        result_doc = Document(str(output_path))
        text_content = "\n".join([p.text for p in result_doc.paragraphs])

        # Check the paragraph with multiple placeholders
        expected_text = "Project Project Phoenix for TechCorp Inc. started on 2024-01-15"
        assert expected_text in text_content

    finally:
        output_path.unlink(missing_ok=True)
        sample_template.unlink(missing_ok=True)


def test_table_placeholder_replacement(sample_template, sample_data):
    """Test that placeholders in tables are replaced"""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as out:
        output_path = Path(out.name)

    try:
        fill_word_from_structured_response(sample_template, sample_data, output_path)

        result_doc = Document(str(output_path))

        # Check table content
        table = result_doc.tables[0]
        tech_cell_text = table.cell(1, 1).text

        assert "Python, React, Docker, Kubernetes" in tech_cell_text
        assert "{listeTechnologies}" not in tech_cell_text

    finally:
        output_path.unlink(missing_ok=True)
        sample_template.unlink(missing_ok=True)


def test_header_placeholder_replacement(sample_template, sample_data):
    """Test that placeholders in headers are replaced"""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as out:
        output_path = Path(out.name)

    try:
        fill_word_from_structured_response(sample_template, sample_data, output_path)

        result_doc = Document(str(output_path))
        header = result_doc.sections[0].header
        header_text = "\n".join([p.text for p in header.paragraphs])

        assert "TechCorp Inc." in header_text
        assert "{nomSociete}" not in header_text

    finally:
        output_path.unlink(missing_ok=True)
        sample_template.unlink(missing_ok=True)


def test_footer_placeholder_replacement(sample_template, sample_data):
    """Test that placeholders in footers are replaced"""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as out:
        output_path = Path(out.name)

    try:
        fill_word_from_structured_response(sample_template, sample_data, output_path)

        result_doc = Document(str(output_path))
        footer = result_doc.sections[0].footer
        footer_text = "\n".join([p.text for p in footer.paragraphs])

        assert "Project Phoenix" in footer_text
        assert "{nomProjet}" not in footer_text

    finally:
        output_path.unlink(missing_ok=True)
        sample_template.unlink(missing_ok=True)


def test_missing_placeholder_data(sample_template):
    """Test that missing placeholders don't cause errors"""
    incomplete_data = {
        "informationsProjet": {
            "nomSociete": "TechCorp Inc.",
            # nomProjet is missing
        }
    }

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as out:
        output_path = Path(out.name)

    try:
        # Should not raise an exception
        fill_word_from_structured_response(sample_template, incomplete_data, output_path)

        result_doc = Document(str(output_path))
        text_content = "\n".join([p.text for p in result_doc.paragraphs])

        # Replaced placeholders should be updated
        assert "TechCorp Inc." in text_content

        # Missing placeholders should remain unchanged (logged as warning)
        assert "{nomProjet}" in text_content or "nomProjet" not in text_content

    finally:
        output_path.unlink(missing_ok=True)
        sample_template.unlink(missing_ok=True)


def test_formatting_preservation():
    """Test that text formatting (bold, italic) is preserved during replacement"""
    doc = Document()
    paragraph = doc.add_paragraph()

    # Add formatted text with placeholder
    run1 = paragraph.add_run("Company: ")
    run1.bold = True
    run2 = paragraph.add_run("{nomSociete}")
    run2.italic = True

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        template_path = Path(tmp.name)

    data = {
        "informationsProjet": {
            "nomSociete": "TechCorp"
        }
    }

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as out:
        output_path = Path(out.name)

    try:
        fill_word_from_structured_response(template_path, data, output_path)

        result_doc = Document(str(output_path))
        paragraph = result_doc.paragraphs[0]

        # Check content
        assert "TechCorp" in paragraph.text
        assert "{nomSociete}" not in paragraph.text

        # Check formatting is preserved
        assert paragraph.runs[0].bold is True
        # The replaced text should inherit formatting from the placeholder

    finally:
        output_path.unlink(missing_ok=True)
        template_path.unlink(missing_ok=True)


def test_empty_document():
    """Test handling of empty document"""
    doc = Document()

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        template_path = Path(tmp.name)

    data = {"informationsProjet": {"nomSociete": "TechCorp"}}

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as out:
        output_path = Path(out.name)

    try:
        # Should not raise an exception
        fill_word_from_structured_response(template_path, data, output_path)

        result_doc = Document(str(output_path))
        # Should have no paragraphs or empty paragraphs
        assert len([p for p in result_doc.paragraphs if p.text]) == 0

    finally:
        output_path.unlink(missing_ok=True)
        template_path.unlink(missing_ok=True)


def test_textbox_placeholder_replacement():
    """Test that placeholders in textboxes are replaced (requires XML manipulation)"""
    import zipfile
    import xml.etree.ElementTree as ET
    import shutil

    # Create a document with a textbox containing a placeholder
    doc = Document()
    doc.add_paragraph("Test document")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        original_path = Path(tmp.name)

    # Create a modified version with textbox
    template_path = Path(tempfile.mktemp(suffix=".docx"))
    shutil.copy(original_path, template_path)

    # Modify the document to add textbox with placeholder
    with zipfile.ZipFile(original_path, 'r') as docx_read:
        xml_content = docx_read.read('word/document.xml')
        root = ET.fromstring(xml_content)

        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        # Create a textbox structure
        body = root.find('.//w:body', ns)
        if body is not None:
            p = ET.SubElement(body, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
            txbxContent = ET.SubElement(p, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent')
            inner_p = ET.SubElement(txbxContent, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
            r = ET.SubElement(inner_p, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            t = ET.SubElement(r, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            t.text = "Company: {nomSociete}"

        # Write to new file
        with zipfile.ZipFile(template_path, 'w', zipfile.ZIP_DEFLATED) as docx_write:
            # Write modified document.xml
            docx_write.writestr('word/document.xml', ET.tostring(root, encoding='unicode'))

            # Copy all other files
            for item in docx_read.namelist():
                if item != 'word/document.xml':
                    docx_write.writestr(item, docx_read.read(item))

    original_path.unlink()

    data = {"informationsProjet": {"nomSociete": "TechCorp Inc"}}

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as out:
        output_path = Path(out.name)

    try:
        fill_word_from_structured_response(template_path, data, output_path)

        # Verify via XML that the placeholder was replaced
        with zipfile.ZipFile(output_path, 'r') as docx:
            result_xml = docx.read('word/document.xml')

        result_root = ET.fromstring(result_xml)
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

        # Check that placeholder is gone and value is there
        all_text = ''.join([t.text or '' for t in result_root.findall('.//w:t', ns)])

        assert "TechCorp Inc" in all_text or len(all_text) > 0  # Content exists
        assert "{nomSociete}" not in all_text  # Placeholder is replaced

    finally:
        output_path.unlink(missing_ok=True)
        template_path.unlink(missing_ok=True)

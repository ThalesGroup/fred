from __future__ import annotations

import io

from docx import Document

from agentic_backend.core.writable_documents.docx_export import markdown_to_docx_bytes


def _read_paragraphs(data: bytes) -> list[tuple[str, str]]:
    doc = Document(io.BytesIO(data))
    return [
        ((p.style.name or "") if p.style is not None else "", p.text)
        for p in doc.paragraphs
    ]


def test_docx_is_valid_zip_with_title():
    data = markdown_to_docx_bytes("# Hello\n\nBody", title="My Doc")
    assert data[:2] == b"PK"  # docx is a zip
    paras = _read_paragraphs(data)
    assert ("Title", "My Doc") in paras
    assert ("Heading 1", "Hello") in paras


def test_docx_maps_common_markdown_constructs():
    md = (
        "# H1\n\n"
        "Some **bold** and *italic* and `code`.\n\n"
        "- bullet one\n"
        "- bullet two\n\n"
        "1. first\n"
        "2. second\n\n"
        "```\n"
        "code line 1\n"
        "code line 2\n"
        "```\n\n"
        "Final soft\n"
        "wrapped paragraph."
    )
    paras = _read_paragraphs(markdown_to_docx_bytes(md))
    styles = [s for s, _ in paras]
    texts = [t for _, t in paras]

    assert ("Heading 1", "H1") in paras
    # Inline markers are stripped; text content preserved in a normal paragraph.
    assert "Some bold and italic and code." in texts
    assert styles.count("List Bullet") == 2
    assert styles.count("List Number") == 2
    # Fenced code block content kept (joined with newlines) in one paragraph.
    assert any("code line 1\ncode line 2" == t for t in texts)
    # Soft-wrapped lines merged into a single paragraph.
    assert "Final soft wrapped paragraph." in texts


def test_docx_handles_empty_content():
    data = markdown_to_docx_bytes("", title="Empty")
    assert data[:2] == b"PK"
    assert ("Title", "Empty") in _read_paragraphs(data)


def test_bold_run_is_marked_bold():
    data = markdown_to_docx_bytes("This is **strong** text")
    doc = Document(io.BytesIO(data))
    runs = [r for p in doc.paragraphs for r in p.runs]
    bold_run = next((r for r in runs if r.text == "strong"), None)
    assert bold_run is not None
    assert bold_run.bold is True


def test_nested_emphasis_does_not_leak_markers():
    # Flat parsing can't render nested/combined emphasis; we keep clean text
    # instead of leaking literal ** or _ characters.
    data = markdown_to_docx_bytes("## **_Combined *Bold* and _Italic_ Heading_**")
    heading = _paragraph_by_style(data, "Heading 2")
    assert "*" not in heading.text
    assert "_" not in heading.text
    assert heading.text == "Combined Bold and Italic Heading"


def test_stray_unmatched_markers_are_stripped():
    data = markdown_to_docx_bytes("A lone * and a stray _ in text")
    doc = Document(io.BytesIO(data))
    text = " ".join(p.text for p in doc.paragraphs)
    assert "*" not in text and "_" not in text


def test_code_span_keeps_literal_markers():
    # Markers inside backticks are literal code and must survive untouched.
    data = markdown_to_docx_bytes("Use `a_b * c` here")
    doc = Document(io.BytesIO(data))
    code_run = next(
        (r for p in doc.paragraphs for r in p.runs if r.text == "a_b * c"), None
    )
    assert code_run is not None


def _paragraph_by_style(data: bytes, style_name: str):
    doc = Document(io.BytesIO(data))
    return next(
        p for p in doc.paragraphs if p.style is not None and p.style.name == style_name
    )


def test_heading_strips_inline_markers_and_marks_bold():
    heading = _paragraph_by_style(
        markdown_to_docx_bytes("## **New section title**"), "Heading 2"
    )
    # Markers are stripped from the visible text...
    assert heading.text == "New section title"
    # ...and the run is actually bold.
    assert heading.runs and all(r.bold for r in heading.runs)


def test_heading_with_partial_italic():
    heading = _paragraph_by_style(
        markdown_to_docx_bytes("# Plan for *Q3*"), "Heading 1"
    )
    assert heading.text == "Plan for Q3"
    italic_run = next((r for r in heading.runs if r.text == "Q3"), None)
    assert italic_run is not None and italic_run.italic is True


def test_title_supports_inline_formatting():
    title = _paragraph_by_style(
        markdown_to_docx_bytes("Body", title="**Bold Title**"), "Title"
    )
    assert title.text == "Bold Title"
    assert title.runs and all(r.bold for r in title.runs)


def _table_grid(data: bytes) -> list[list[str]]:
    doc = Document(io.BytesIO(data))
    assert doc.tables, "expected at least one table"
    table = doc.tables[0]
    return [[cell.text for cell in row.cells] for row in table.rows]


def test_table_is_rendered_with_header_and_rows():
    md = (
        "| Name  | Role     |\n"
        "|-------|----------|\n"
        "| Alice | Engineer |\n"
        "| Bob   | Designer |\n"
    )
    data = markdown_to_docx_bytes(md)
    grid = _table_grid(data)
    assert grid == [
        ["Name", "Role"],
        ["Alice", "Engineer"],
        ["Bob", "Designer"],
    ]


def test_table_header_cells_are_bold():
    md = "| A | B |\n|---|---|\n| 1 | 2 |\n"
    doc = Document(io.BytesIO(markdown_to_docx_bytes(md)))
    header_cells = doc.tables[0].rows[0].cells
    for cell in header_cells:
        runs = [r for p in cell.paragraphs for r in p.runs]
        assert runs and all(r.bold for r in runs)


def test_table_ragged_rows_are_padded():
    md = "| A | B | C |\n|---|---|---|\n| 1 | 2 |\n"
    grid = _table_grid(markdown_to_docx_bytes(md))
    assert grid == [["A", "B", "C"], ["1", "2", ""]]


def test_table_supports_inline_formatting_in_cells():
    md = "| Col |\n|-----|\n| **bold** |\n"
    doc = Document(io.BytesIO(markdown_to_docx_bytes(md)))
    body_cell = doc.tables[0].rows[1].cells[0]
    runs = [r for p in body_cell.paragraphs for r in p.runs]
    bold_run = next((r for r in runs if r.text == "bold"), None)
    assert bold_run is not None and bold_run.bold is True


def test_table_honors_escaped_pipe():
    md = "| Expr |\n|------|\n| a \\| b |\n"
    grid = _table_grid(markdown_to_docx_bytes(md))
    assert grid == [["Expr"], ["a | b"]]


def test_pipe_without_separator_is_not_a_table():
    md = "This sentence | has a pipe but no separator.\n"
    doc = Document(io.BytesIO(markdown_to_docx_bytes(md)))
    assert not doc.tables
    assert any("has a pipe" in p.text for p in doc.paragraphs)

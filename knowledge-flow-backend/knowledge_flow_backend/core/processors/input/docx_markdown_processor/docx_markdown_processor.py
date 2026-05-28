# Copyright Thales 2025
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import base64
import logging
import re
import subprocess
import xml.etree.ElementTree as ET
import zipfile
from datetime import datetime
from pathlib import Path
from shutil import which

from docx import Document

from knowledge_flow_backend.application_context import get_configuration
from knowledge_flow_backend.common.processing_profile_context import get_current_processing_profile
from knowledge_flow_backend.core.processors.input.common.base_input_processor import BaseMarkdownProcessor, InputConversionError
from knowledge_flow_backend.core.processors.input.common.image_describer import build_image_describer

logger = logging.getLogger(__name__)

_RASTER_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".tif", ".tiff"}


def _replace_image_reference(md_content: str, img_path: str, description: str) -> str:
    """
    Replace an image reference with its description text, handling all pandoc output styles:
    - Raw HTML:        <img src="path" ... />  (markdown_strict with sized images)
    - Inline:          ![alt](path)
    - Reference-style: ![alt][label] + [label]: path  (produced by --reference-links)
    """
    replacement = f"\n\n{description}\n\n"
    safe_replacement = replacement.replace("\\", r"\\")
    path_re = re.escape(img_path)

    # Raw HTML img tag: <img src="path" ... /> — markdown_strict emits this for sized images
    new = re.sub(r'<img\s[^>]*src="' + path_re + r'"[^>]*/>', safe_replacement, md_content)
    if new != md_content:
        return new

    # Inline link: ![alt](path) or ![alt](path "title")
    new = re.sub(r"!\[[^\]]*\]\(" + path_re + r'(?:\s+"[^"]*")?\)', safe_replacement, md_content)
    if new != md_content:
        return new

    # Reference-style: locate the definition [label]: path, then replace usages
    ref_def = re.search(r"^\[([^\]]+)\]:\s*" + path_re + r"[^\n]*$", md_content, re.MULTILINE)
    if ref_def:
        label = ref_def.group(1)
        label_re = re.escape(label)
        # Replace ![alt][label] and shorthand ![label]
        new = re.sub(r"!\[[^\]]*\]\[" + label_re + r"\]", safe_replacement, md_content)
        new = re.sub(r"!\[" + label_re + r"\](?!\[)", safe_replacement, new)
        # Remove the reference definition line
        new = re.sub(r"^\[" + label_re + r"\]:\s*" + path_re + r"[^\n]*\n?", "", new, flags=re.MULTILINE)
        return new

    return md_content


def default_or_unknown(value: str, default="None") -> str:
    return value.strip() if value and value.strip() else default


class DocxMarkdownProcessor(BaseMarkdownProcessor):
    description = "Converts DOCX files to Markdown while preserving headings, tables, and basic formatting."

    def __init__(self):
        super().__init__()
        self.image_describer = None
        self._warned_missing_vision_model = False

    def _resolve_image_describer(self, process_images: bool):
        if not process_images:
            return None
        if self.image_describer is not None:
            return self.image_describer
        if not get_configuration().vision_model:
            if not self._warned_missing_vision_model:
                logger.warning("[PROCESSOR][DOCX] Vision model configuration is missing while process_images is enabled.")
                self._warned_missing_vision_model = True
            return None
        self.image_describer = build_image_describer(get_configuration().vision_model)
        return self.image_describer

    def _annotate_markdown_tables(self, md_content: str) -> str:
        """Wrap Markdown tables with TABLE_START/END markers for downstream chunking."""

        def is_pipe_separator(line: str) -> bool:
            s = line.strip()
            if "|" not in s or "-" not in s:
                return False
            return all(ch in "|-: " for ch in s)

        def is_pipe_row(line: str) -> bool:
            s = line.strip()
            return "|" in s and not s.startswith("```")

        def is_grid_border(line: str) -> bool:
            s = line.strip()
            return s.startswith("+") and s.endswith("+") and any(ch in s for ch in "-=")

        def is_grid_row(line: str) -> bool:
            s = line.strip()
            return s.startswith("|") and s.endswith("|")

        lines = md_content.splitlines()
        out: list[str] = []
        i = 0
        table_idx = 0
        while i < len(lines):
            line = lines[i]

            # Preserve existing annotations untouched
            if line.strip().startswith("<!-- TABLE_START"):
                out.append(line)
                i += 1
                while i < len(lines):
                    out.append(lines[i])
                    if lines[i].strip().startswith("<!-- TABLE_END"):
                        i += 1
                        break
                    i += 1
                continue

            # Pipe table detection: header + separator
            if i + 1 < len(lines) and is_pipe_row(line) and is_pipe_separator(lines[i + 1]):
                start = i
                i += 2
                while i < len(lines) and lines[i].strip() and is_pipe_row(lines[i]):
                    i += 1
                table_md = "\n".join(lines[start:i]).strip()
                if table_md:
                    table_idx += 1
                    out.append(f"<!-- TABLE_START:id=docx_{table_idx} -->")
                    out.append(table_md)
                    out.append("<!-- TABLE_END -->")
                continue

            # Grid table detection (pandoc-style)
            if is_grid_border(line):
                start = i
                i += 1
                while i < len(lines) and lines[i].strip() and (is_grid_border(lines[i]) or is_grid_row(lines[i])):
                    i += 1
                table_md = "\n".join(lines[start:i]).strip()
                if table_md:
                    table_idx += 1
                    out.append(f"<!-- TABLE_START:id=docx_{table_idx} -->")
                    out.append(table_md)
                    out.append("<!-- TABLE_END -->")
                continue

            out.append(line)
            i += 1

        # Preserve trailing newline behavior
        return "\n".join(out)

    def check_file_validity(self, file_path: Path) -> bool:
        try:
            with zipfile.ZipFile(file_path, "r") as docx_zip:
                return "word/document.xml" in docx_zip.namelist()
        except zipfile.BadZipFile:
            logger.error(f"{file_path} n'est pas une archive ZIP valide.")
        except Exception as e:
            logger.error(f"Erreur inattendue lors de la vérification de {file_path}: {e}")
        return False

    def _collect_non_empty_paragraphs(self, paragraphs, *, limit: int | None = None) -> list[str]:
        lines: list[str] = []
        for paragraph in paragraphs:
            text = (getattr(paragraph, "text", "") or "").strip()
            if not text:
                continue
            lines.append(text)
            if limit is not None and len(lines) >= limit:
                break
        return lines

    def _extract_text_from_docx_xml_part(self, docx_zip: zipfile.ZipFile, part_name: str) -> list[str]:
        try:
            xml_bytes = docx_zip.read(part_name)
            root = ET.fromstring(xml_bytes)
        except Exception as e:
            logger.warning("[DOCX][GUARDRAIL] Failed to parse %s: %s", part_name, e)
            return []

        lines: list[str] = []
        for elem in root.iter():
            if elem.tag.endswith("}t") and elem.text:
                text = elem.text.strip()
                if text:
                    lines.append(text)

        return lines

    def _extract_header_footer_xml_lines(self, file_path: Path) -> tuple[list[str], list[str]]:
        header_lines: list[str] = []
        footer_lines: list[str] = []

        try:
            with zipfile.ZipFile(file_path, "r") as docx_zip:
                for part_name in docx_zip.namelist():
                    if part_name.startswith("word/header") and part_name.endswith(".xml"):
                        header_lines.extend(self._extract_text_from_docx_xml_part(docx_zip, part_name))
                    elif part_name.startswith("word/footer") and part_name.endswith(".xml"):
                        footer_lines.extend(self._extract_text_from_docx_xml_part(docx_zip, part_name))
        except Exception as e:
            logger.warning("[DOCX][GUARDRAIL] Failed to inspect DOCX xml parts for %s: %s", file_path, e)

        return header_lines, footer_lines

    def extract_guardrail_text(self, file_path: Path) -> str | None:
        """
        Extract lightweight boundary text for ingestion guardrails.

        Why this exists:
        - DOCX files often carry visible classification markings in headers
          and footers.
        - Some Office/label integrations store visible footer/header markings
          inside drawing/textbox XML, which python-docx paragraph APIs may not
          expose directly.

        How to use:
        - Collect non-empty header/footer paragraphs from every section.
        - Also inspect raw DOCX header/footer XML parts as a fallback for
          textbox/shape-based markings.
        - Add a small body head/tail fallback in case the visible marking is
          rendered in the main document flow.
        """
        try:
            doc = Document(str(file_path))
        except Exception as e:
            logger.warning("[DOCX][GUARDRAIL] Failed to open %s: %s", file_path, e)
            return None

        header_lines: list[str] = []
        footer_lines: list[str] = []

        for section in doc.sections:
            header_lines.extend(self._collect_non_empty_paragraphs(section.header.paragraphs, limit=10))
            footer_lines.extend(self._collect_non_empty_paragraphs(section.footer.paragraphs, limit=10))

        header_xml_lines, footer_xml_lines = self._extract_header_footer_xml_lines(file_path)

        body_lines = self._collect_non_empty_paragraphs(doc.paragraphs)
        body_head = body_lines[:5]
        body_tail = body_lines[-5:] if len(body_lines) > 5 else []

        parts: list[str] = []

        if header_lines:
            parts.append("HEADER:")
            parts.extend(dict.fromkeys(header_lines))

        if footer_lines:
            if parts:
                parts.append("")
            parts.append("FOOTER:")
            parts.extend(dict.fromkeys(footer_lines))

        if header_xml_lines:
            if parts:
                parts.append("")
            parts.append("HEADER_XML:")
            parts.extend(dict.fromkeys(header_xml_lines))

        if footer_xml_lines:
            if parts:
                parts.append("")
            parts.append("FOOTER_XML:")
            parts.extend(dict.fromkeys(footer_xml_lines))

        if body_head:
            if parts:
                parts.append("")
            parts.append("BODY_HEAD:")
            parts.extend(body_head)

        if body_tail:
            if parts:
                parts.append("")
            parts.append("BODY_TAIL:")
            parts.extend(body_tail)

        result = "\n".join(parts).strip()
        return result or None

    def extract_file_metadata(self, file_path: Path) -> dict:
        try:
            doc = Document(str(file_path))
            cp = doc.core_properties
            return {
                # identity
                "title": cp.title or None,
                "author": cp.author or None,
                "created": cp.created if isinstance(cp.created, datetime) else None,
                "modified": cp.modified if isinstance(cp.modified, datetime) else None,
                "last_modified_by": cp.last_modified_by or None,
                # optional extras (kept out of vector index; good for UI/analytics)
                "extras": {
                    "docx.core.category": cp.category or None,
                    "docx.core.subject": cp.subject or None,
                    "docx.core.keywords": cp.keywords or None,
                },
            }
        except Exception as e:
            logger.error(f"Error extracting metadata for {file_path}: {e}")
            return {"document_name": file_path.name, "error": str(e)}

    def convert_file_to_markdown(self, file_path: Path, output_dir: Path, document_uid: str | None) -> dict:
        output_dir.mkdir(parents=True, exist_ok=True)
        md_path = output_dir / "output.md"

        processing = get_configuration().processing
        current_profile = get_current_processing_profile()
        active_profile = processing.normalize_profile(current_profile)
        profile_cfg = processing.get_profile_config(active_profile)
        image_describer = self._resolve_image_describer(profile_cfg.process_images)

        images_dir = output_dir
        extra_args = [f"--extract-media={images_dir}", "--preserve-tabs", "--wrap=none", "--reference-links"]

        try:
            subprocess.run(
                [
                    "pandoc",
                    "--to",
                    "markdown_strict+pipe_tables",
                    str(file_path),
                    "-o",
                    str(md_path),
                    *extra_args,
                ],
                check=True,
                capture_output=True,
                text=True,
            )
        except subprocess.CalledProcessError as exc:
            stderr = (exc.stderr or "").strip()
            stdout = (exc.stdout or "").strip()
            detail = stderr or stdout or str(exc)
            raise InputConversionError(f"Pandoc DOCX conversion failed for '{file_path.name}': {detail}") from exc

        # pypandoc.convert_file(str(file_path), to="markdown_strict+pipe_tables", outputfile=str(md_path), extra_args=extra_args)

        # Convert EMF to SVG
        for img_path in (images_dir / "media").glob("*.emf"):
            svg_path = img_path.with_suffix(".svg")
            if which("inkscape") is None:
                logger.error("[DOCX] Inkscape not found; cannot convert %s to SVG. Leaving EMF in place.", img_path)
                continue

            try:
                subprocess.run(["inkscape", str(img_path), "--export-filename=" + str(svg_path)], check=True)
            except subprocess.CalledProcessError as e:
                logger.error("[DOCX] Inkscape failed converting %s to SVG: %s", img_path, e)
                continue

            # Remove the original EMF file
            img_path.unlink()

        # Update references in the markdown file
        with open(md_path, "r", encoding="utf-8") as f:
            md_content = f.read()

        md_content = md_content.replace(".emf", ".svg")

        # Describe raster images with the vision model when enabled
        if image_describer is not None:
            media_dir = output_dir / "media"
            if media_dir.is_dir():
                for img_path in sorted(media_dir.iterdir()):
                    if img_path.suffix.lower() not in _RASTER_IMAGE_SUFFIXES:
                        continue
                    try:
                        img_b64 = base64.b64encode(img_path.read_bytes()).decode("utf-8")
                        description = image_describer.describe(img_b64)
                        logger.info("[PROCESSOR][DOCX] Described image: %s", img_path.name)
                    except Exception as e:
                        logger.warning("[PROCESSOR][DOCX] Image description failed for %s: %s", img_path.name, e)
                        description = "Image could not be described."
                    before = md_content
                    md_content = _replace_image_reference(md_content, str(img_path), description)
                    if md_content == before:
                        logger.warning("[PROCESSOR][DOCX] No match found for image reference: %s — regex did not replace", img_path)

        # Change media path to use api endpoint
        md_content = md_content.replace(str(output_dir), f"knowledge-flow/v1/markdown/{document_uid}")

        # Annotate tables so chunker keeps them intact
        md_content = self._annotate_markdown_tables(md_content)

        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_content)

        for f in output_dir.glob("*.lua"):
            f.unlink()

        return {"doc_dir": str(output_dir), "md_file": str(md_path)}

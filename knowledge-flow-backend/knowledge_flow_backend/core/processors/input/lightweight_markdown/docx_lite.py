from __future__ import annotations

import logging
from pathlib import Path
from typing import List

from docx import Document

from .lite_types import (
    LiteMarkdownOptions,
    LiteMarkdownResult,
    LitePageMarkdown,
    collapse_whitespace,
    enforce_max_chars,
)

logger = logging.getLogger(__name__)


def _heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    name = style_name.strip().lower()
    if name.startswith("heading "):
        try:
            lvl = int(name.split(" ", 1)[1])
            return max(1, min(lvl, 6))
        except Exception:
            return None
    return None


def _is_list_style(style_name: str | None) -> bool:
    if not style_name:
        return False
    name = style_name.strip().lower()
    return "list" in name or "bullet" in name or "number" in name


def _cell_text(cell) -> str:
    # Join paragraphs within a cell; ignore internal tables for lightweight pass
    return "\n".join(p.text.strip() for p in cell.paragraphs if p.text and p.text.strip())


class DocxLiteMarkdownExtractor:
    """Fast DOCX → Markdown extraction using python-docx.

    - Preserves headings via paragraph styles
    - Emits simple bullets for list-like styles
    - Converts tables into small pipe tables, truncated by options
    - Skips images and complex formatting
    - Returns in-memory strings only
    """

    def extract(self, file_path: Path, options: LiteMarkdownOptions | None = None) -> LiteMarkdownResult:
        opts = options or LiteMarkdownOptions()
        doc = Document(str(file_path))

        lines: List[str] = []

        # paragraphs first
        for p in doc.paragraphs:
            text = p.text or ""
            text = text.strip()
            if not text:
                continue

            lvl = _heading_level(getattr(p.style, "name", None))
            if lvl:
                lines.append(f"{'#' * lvl} {text}")
                continue

            if _is_list_style(getattr(p.style, "name", None)):
                lines.append(f"- {text}")
                continue

            lines.append(text)

        # tables (lightweight)
        if opts.include_tables and getattr(doc, "tables", None):
            table_idx = 0
            for t in doc.tables:
                table_idx += 1
                try:
                    rows = min(len(t.rows), max(0, opts.max_table_rows))
                    cols = min(len(t.columns), max(0, opts.max_table_cols))
                    if rows == 0 or cols == 0:
                        continue

                    lines.append("")
                    lines.append(f"<!-- TABLE_START:index={table_idx} rows={len(t.rows)} cols={len(t.columns)} -->")

                    # header from first row
                    header = [ _cell_text(t.cell(0, c)) for c in range(cols) ]
                    header = [ h if h else " " for h in header ]
                    lines.append("| " + " | ".join(header) + " |")
                    lines.append("| " + " | ".join(["---"] * cols) + " |")

                    for r in range(1, rows):
                        row_cells = [ _cell_text(t.cell(r, c)) for c in range(cols) ]
                        row_cells = [ c if c else " " for c in row_cells ]
                        lines.append("| " + " | ".join(row_cells) + " |")

                    if rows < len(t.rows) or cols < len(t.columns):
                        lines.append(f"… (table truncated)")

                    lines.append("<!-- TABLE_END -->")
                    lines.append("")
                except Exception as e:
                    logger.debug(f"Table {table_idx} export skipped: {e}")

        md = "\n".join(lines)
        if opts.normalize_whitespace:
            md = collapse_whitespace(md)

        md, truncated = enforce_max_chars(md, opts.max_chars)

        result = LiteMarkdownResult(
            document_name=file_path.name,
            page_count=None,  # DOCX has no reliable page model here
            total_chars=len(md),
            truncated=truncated,
            markdown=md,
            pages=[LitePageMarkdown(page_no=1, markdown=md, char_count=len(md))] if opts.return_per_page else [],
            extras={},
        )
        return result


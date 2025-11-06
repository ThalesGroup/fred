# Copyright Thales 2025
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

from __future__ import annotations

import logging
from pathlib import Path
from typing import List

from docx import Document

from knowledge_flow_backend.core.processors.input.common.base_input_processor import (
    BaseMarkdownProcessor,
)

from .lite_types import (
    LiteMarkdownOptions,
    LiteMarkdownResult,
    LitePageMarkdown,
    collapse_whitespace,
    enforce_max_chars,
)

logger = logging.getLogger(__name__)


def _heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    name = style_name.strip().lower()
    if name.startswith("heading "):
        try:
            lvl = int(name.split(" ", 1)[1])
            return max(1, min(lvl, 6))
        except Exception:
            return None
    return None


def _is_list_style(style_name: str | None) -> bool:
    if not style_name:
        return False
    name = style_name.strip().lower()
    return "list" in name or "bullet" in name or "number" in name


def _cell_text(cell) -> str:
    # Join paragraphs within a cell; ignore internal tables for lightweight pass
    return "\n".join(p.text.strip() for p in cell.paragraphs if p.text and p.text.strip())


class DocxLiteMarkdownExtractor:
    """Fast DOCX → Markdown extraction using python-docx.

    - Preserves headings via paragraph styles
    - Emits simple bullets for list-like styles
    - Converts tables into small pipe tables, truncated by options
    - Skips images and complex formatting
    - Returns in-memory strings only
    """

    def extract(self, file_path: Path, options: LiteMarkdownOptions | None = None) -> LiteMarkdownResult:
        opts = options or LiteMarkdownOptions()
        doc = Document(str(file_path))

        lines: List[str] = []

        # paragraphs first
        for p in doc.paragraphs:
            text = p.text or ""
            text = text.strip()
            if not text:
                continue

            lvl = _heading_level(getattr(p.style, "name", None))
            if lvl:
                lines.append(f"{'#' * lvl} {text}")
                continue

            if _is_list_style(getattr(p.style, "name", None)):
                lines.append(f"- {text}")
                continue

            lines.append(text)

        # tables (lightweight)
        if opts.include_tables and getattr(doc, "tables", None):
            table_idx = 0
            for t in doc.tables:
                table_idx += 1
                try:
                    rows = min(len(t.rows), max(0, opts.max_table_rows))
                    cols = min(len(t.columns), max(0, opts.max_table_cols))
                    if rows == 0 or cols == 0:
                        continue

                    lines.append("")
                    lines.append(f"<!-- TABLE_START:index={table_idx} rows={len(t.rows)} cols={len(t.columns)} -->")

                    # header from first row
                    header = [_cell_text(t.cell(0, c)) for c in range(cols)]
                    header = [h if h else " " for h in header]
                    lines.append("| " + " | ".join(header) + " |")
                    lines.append("| " + " | ".join(["---"] * cols) + " |")

                    for r in range(1, rows):
                        row_cells = [_cell_text(t.cell(r, c)) for c in range(cols)]
                        row_cells = [c if c else " " for c in row_cells]
                        lines.append("| " + " | ".join(row_cells) + " |")

                    if rows < len(t.rows) or cols < len(t.columns):
                        lines.append("… (table truncated)")

                    lines.append("<!-- TABLE_END -->")
                    lines.append("")
                except Exception as e:
                    logger.debug(f"Table {table_idx} export skipped: {e}")

        md = "\n".join(lines)
        if opts.normalize_whitespace:
            md = collapse_whitespace(md)

        md, truncated = enforce_max_chars(md, opts.max_chars)

        result = LiteMarkdownResult(
            document_name=file_path.name,
            page_count=None,  # DOCX has no reliable page model here
            total_chars=len(md),
            truncated=truncated,
            markdown=md,
            pages=[LitePageMarkdown(page_no=1, markdown=md, char_count=len(md))] if opts.return_per_page else [],
            extras={},
        )
        return result


class DocxLiteMarkdownProcessor(BaseMarkdownProcessor):
    """BaseMarkdownProcessor-compatible wrapper around the lightweight DOCX extractor.

    Keeps behavior homogeneous with other processors by exposing the
    check/metadata/convert interface and always writing to `output.md` in the
    provided `output_dir`.
    """

    def __init__(self) -> None:
        super().__init__()
        self._extractor = DocxLiteMarkdownExtractor()

    def check_file_validity(self, file_path: Path) -> bool:
        try:
            # python-docx will raise if the file isn't a valid DOCX
            Document(str(file_path))
            return True
        except Exception as e:
            logger.error(f"Invalid or corrupted DOCX file: {file_path} - {e}")
            return False

    def extract_file_metadata(self, file_path: Path) -> dict:
        # Keep lightweight and consistent with other simple processors
        try:
            size = file_path.stat().st_size
        except Exception:
            size = None
        return {
            "document_name": file_path.name,
            "file_size_bytes": size,
            "suffix": file_path.suffix,
        }

    def convert_file_to_markdown(self, file_path: Path, output_dir: Path, document_uid: str | None) -> dict:
        output_dir.mkdir(parents=True, exist_ok=True)
        md_path = output_dir / "output.md"

        # Use defaults; whitespace normalization is handled by the extractor
        opts = LiteMarkdownOptions()
        result = self._extractor.extract(file_path, options=opts)
        md_path.write_text(result.markdown, encoding="utf-8")

        return {"doc_dir": str(output_dir), "md_file": str(md_path)}
